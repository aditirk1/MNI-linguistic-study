"""
Build a structured Word report with embedded figures.

Generates new figures under report_assets/generated/ and compiles
NEBULA_Multimodal_Report.docx in the project root.

Run:
  .venv\\Scripts\\python.exe build_project_report_docx.py
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).resolve().parent
ASSETS = BASE / "report_assets"
GEN = ASSETS / "generated"
DOCX_OUT = BASE / "NEBULA_Multimodal_Report.docx"

ROI_NAMES = ["IFG_L", "STG_L", "AngG_L", "SMA_L", "IFG_R", "STG_R"]


def p(txt: str) -> str:
    return " ".join(txt.split())


def ensure_assets() -> None:
    GEN.mkdir(parents=True, exist_ok=True)


def fig_mean_fc_matrix() -> Path:
    """Group-mean Fisher-z 6x6 from final nilearn_fc_motion pipeline."""
    out = GEN / "fig_group_mean_fc_matrix.png"
    npys = sorted((BASE / "derivatives/nilearn_fc_motion").glob("sub-*_fc_matrix.npy"))
    if not npys:
        raise SystemExit("No sub-*_fc_matrix.npy under derivatives/nilearn_fc_motion/")
    mats = [np.load(f) for f in npys]
    mean_z = np.mean(np.stack(mats, axis=0), axis=0)
    n_sub = len(npys)

    off = mean_z[np.triu_indices(6, k=1)]
    vmax = float(np.max(np.abs(off)))
    vmax = max(vmax, 0.15)

    mask = np.eye(6, dtype=bool)
    fig, ax = plt.subplots(figsize=(7.2, 6.4))
    sns.heatmap(
        mean_z,
        mask=mask,
        annot=True,
        fmt=".2f",
        cmap="RdBu_r",
        center=0.0,
        vmin=-vmax,
        vmax=vmax,
        square=True,
        linewidths=0.5,
        linecolor="white",
        xticklabels=ROI_NAMES,
        yticklabels=ROI_NAMES,
        cbar_kws={"label": "Fisher z", "shrink": 0.85},
        ax=ax,
    )
    ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha="right")
    ax.set_title(
        f"Group-mean Fisher-z resting FC (n={n_sub}, MNI BOLD, motion-regressed)\n"
        f"diagonal omitted; off-diagonal range [{off.min():.2f}, {off.max():.2f}]",
        fontsize=11,
    )
    fig.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight", pad_inches=0.35)
    plt.close()
    return out


def fig_cohort_predictors() -> Path:
    out = GEN / "fig_cohort_predictors.png"
    subs = [
        s.strip()
        for s in (BASE / "subset_50_participants.txt").read_text().splitlines()
        if s.strip()
    ]
    dm = pd.read_csv(BASE / "shared_design_matrix.csv")
    d = dm[dm["participant_id"].isin(subs)].copy()

    fig, axes = plt.subplots(1, 2, figsize=(9, 4))
    ax = axes[0]
    vc = d["nlang"].value_counts().sort_index()
    ax.bar(vc.index.astype(str), vc.values, color="steelblue", edgecolor="black")
    ax.set_xlabel("Number of languages (nlang)")
    ax.set_ylabel("Count")
    ax.set_title("Language count in analysis cohort (n=51)")

    ax = axes[1]
    ax.hist(d["entropy_curr_tot_exp"], bins=12, color="coral", edgecolor="black")
    ax.set_xlabel("Entropy of exposure (entropy_curr_tot_exp)")
    ax.set_ylabel("Count")
    ax.set_title("Distribution of language-exposure balance")
    fig.suptitle("Self-report predictors (LEAP-Q derived)", y=1.02)
    plt.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close()
    return out


def fig_motion_fd() -> Path:
    out = GEN / "fig_mean_fd_distribution.png"
    fc = pd.read_csv(BASE / "derivatives/nilearn_fc_motion/fc_outcomes_motion.csv")
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.hist(fc["mean_FD"], bins=15, color="seagreen", edgecolor="black")
    ax.axvline(fc["mean_FD"].mean(), color="darkred", linestyle="--", label=f"Mean = {fc['mean_FD'].mean():.3f} mm")
    ax.set_xlabel("Mean framewise displacement (Power et al., mm)")
    ax.set_ylabel("Participants")
    ax.set_title("Head motion summary (resting-state run, n=51)")
    ax.legend()
    plt.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close()
    return out


def fig_predictor_correlation() -> Path:
    out = GEN / "fig_nlang_entropy_scatter.png"
    subs = [
        s.strip()
        for s in (BASE / "subset_50_participants.txt").read_text().splitlines()
        if s.strip()
    ]
    dm = pd.read_csv(BASE / "shared_design_matrix.csv")
    d = dm[dm["participant_id"].isin(subs)]
    fig, ax = plt.subplots(figsize=(5.5, 5))
    ax.scatter(
        d["nlang"],
        d["entropy_curr_tot_exp"],
        alpha=0.75,
        c="navy",
        edgecolors="white",
        s=55,
    )
    r = d["nlang"].corr(d["entropy_curr_tot_exp"])
    ax.set_xlabel("nlang (language count)")
    ax.set_ylabel("Entropy of language exposure")
    ax.set_title(f"nlang vs entropy (r = {r:.2f})")
    plt.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close()
    return out


def fig_fc_vs_nlang() -> Path:
    """Descriptive: mean network FC vs nlang (not adjusted for other covariates)."""
    out = GEN / "fig_mean_fc_vs_nlang.png"
    fc = pd.read_csv(BASE / "derivatives/nilearn_fc_motion/fc_outcomes_motion.csv")
    dm = pd.read_csv(BASE / "shared_design_matrix.csv")
    m = fc.merge(dm, on="participant_id")
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.scatter(m["nlang"], m["mean_lang_fc"], alpha=0.7, c="purple", edgecolors="k", s=45)
    ax.set_xlabel("nlang")
    ax.set_ylabel("Mean Fisher-z (15 edges)")
    ax.set_title("Descriptive: mean resting FC vs language count")
    plt.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close()
    return out


def add_picture_safe(doc: Document, path: Path, width_in: float = 6.0) -> bool:
    if not path.is_file():
        return False
    try:
        doc.add_picture(str(path), width=Inches(width_in))
        return True
    except Exception:
        doc.add_paragraph(f"[Could not embed: {path.name}]")
        return False


def add_figure_caption(doc: Document, text: str) -> None:
    """Short caption above a figure (normal paragraph, not Word list style)."""
    p = doc.add_paragraph()
    r0 = p.add_run("Figure. ")
    r0.bold = True
    p.add_run(text)


def build_document(generated: list[Path], extra_figures: list[tuple[Path, str]]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Multilingual experience and multimodal brain MRI")
    run.bold = True
    run.font.size = Pt(18)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run("NEBULA101 (OpenNeuro ds005613) — analysis report").italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "[YOUR NAME]  |  [COURSE / PROGRAM]  |  [DATE]\n"
        "[ADVISOR or instructor if required]"
    )

    doc.add_paragraph()

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        p(
            """This report describes a multimodal neuroimaging analysis of language diversity and
            balanced language exposure in healthy young adults from the public NEBULA101 dataset.
            Self-report measures included the number of languages spoken (nlang) and Shannon entropy
            of daily language exposure. Resting-state functional MRI was preprocessed with SPM12/25
            (segmentation, coregistration, normalisation to MNI, smoothing) and summarised with
            seed-based functional connectivity (Nilearn) including motion regression. Diffusion MRI
            was preprocessed with FSL (eddy, tensor fit) and analysed with tract-based spatial
            statistics (TBSS) and permutation inference (randomise, TFCE). In n=51 participants with
            complete T1-weighted, resting fMRI, and DWI data, neither nlang nor entropy showed
            significant associations with the pre-specified functional connectivity summaries after
            Bonferroni adjustment across three outcomes, nor with fractional anisotropy on the
            white-matter skeleton at family-wise error p < 0.05. Sex and head motion covariates
            accounted for appreciable variance in connectivity. Results constrain the magnitude of
            linear questionnaire–brain associations in this subsample and are discussed with respect
            to power, construct overlap, and design limitations."""
        )
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        p(
            """A large fraction of the world's population uses more than one language in daily life.
            Cognitive neuroscience has long asked whether multilingual experience is associated with
            differences in brain structure and function. Findings are mixed: some studies report
            altered white-matter microstructure or functional connectivity in multilinguals, while
            others find small or inconsistent effects. One source of heterogeneity is how
            “multilingualism” is operationalised. Counting languages captures diversity of
            repertoires; entropy of self-reported exposure captures how evenly languages are used.
            These constructs are related but not identical, and both can be derived from the
            Language Experience and Proficiency Questionnaire (LEAP-Q) bundled with NEBULA101."""
        )
    )
    doc.add_paragraph(
        p(
            """The present project uses OpenNeuro dataset ds005613 (NEBULA101), which provides
            multimodal MRI and rich language-history data for 101 adults. For feasibility, a subset
            of n=51 participants was analysed with all three modalities locally available,
            stratified to preserve variation in nlang."""
        )
    )

    doc.add_heading("2. Research questions and hypotheses", level=1)
    doc.add_paragraph(
        p(
            """The primary research question was whether nlang and entropy make independent
            contributions to (a) resting-state functional connectivity among a priori language-related
            seeds in MNI space and (b) voxelwise fractional anisotropy on the TBSS skeleton, after
            adjusting for age, education, sex, and (for fMRI) mean head motion."""
        )
    )
    doc.add_paragraph(
        p(
            """Two directional hypotheses were evaluated at the group level. The structural
            hypothesis was that higher nlang and/or higher entropy would be associated with higher
            fractional anisotropy along language-related white-matter trajectories on the TBSS
            skeleton. The functional hypothesis was that the same predictors would be associated with
            altered resting-state coupling within the seeded language circuit (direction left open,
            because both increased segregation and “efficiency”-like decreases have been discussed
            in the literature)."""
        )
    )

    doc.add_heading("3. Methods", level=1)

    doc.add_heading("3.1 Participants and design", level=2)
    doc.add_paragraph(
        p(
            """The analysis cohort comprised n=51 participants listed in subset_50_participants.txt,
            drawn from the 101-subject NEBULA101 release. Demographics and z-scored predictors for
            second-level models come from shared_design_matrix.csv. Pearson correlation between
            nlang and entropy in this cohort was approximately r = 0.43, so unique (partial) effects
            should be interpreted as modestly constrained by shared variance."""
        )
    )
    add_figure_caption(
        doc,
        "Distribution of nlang and of exposure entropy in the n=51 cohort (generated for this report).",
    )
    add_picture_safe(doc, GEN / "fig_cohort_predictors.png", 6.2)
    add_figure_caption(
        doc,
        "Scatter of nlang versus entropy; predictors are correlated and partial effects should be interpreted accordingly.",
    )
    add_picture_safe(doc, GEN / "fig_nlang_entropy_scatter.png", 4.8)

    doc.add_heading("3.2 Imaging data", level=2)
    doc.add_paragraph(
        p(
            """MRI was acquired on a Siemens MAGNETOM Prisma 3 T (Campus Biotech Geneva). Resting
            fMRI used a multiband EPI sequence (TR = 2 s, whole-brain coverage). High-resolution
            T1-weighted anatomical images were used for segmentation and MNI warps. Diffusion data
            were multi-shell (including b ≈ 0, 700, 1000, 2800 s/mm²) with many gradient directions."""
        )
    )

    doc.add_heading("3.3 Resting fMRI preprocessing", level=2)
    doc.add_paragraph(
        p(
            """Six-degree-of-freedom motion parameters were estimated with SPM realign (estimate).
            For spatial normalisation to MNI152, each subject’s mean resting BOLD volume was
            coregistered to the anatomical T1; unified segmentation of the T1 supplied a forward
            deformation field (y_*) with which the time series (all frames) were written into MNI
            space (3 mm isotropic), then smoothed (6 mm FWHM Gaussian). Smoothed warped BOLD
            (sw*.nii) was the input to connectivity analyses. Full batch details and logs are in
            normalize_bold_to_mni.m and normalize_bold_log.txt."""
        )
    )

    doc.add_heading("3.4 Functional connectivity and second-level model", level=2)
    doc.add_paragraph(
        p(
            """Six spherical seeds (8 mm radius) in MNI space were defined in fc_with_motion.py
            (IFG_L, STG_L, AngG_L, SMA_L, IFG_R, STG_R). Time series were band-pass filtered
            (0.01–0.1 Hz), detrended, and z-scored; the six SPM motion parameters were regressed
            from each seed. Subject-level Fisher-z connectivity matrices were saved as npy files.
            Three outcomes entered group OLS models in Python (statsmodels): mean Fisher-z across
            all 15 unique edges, left IFG–STG z, and right IFG–STG z. Predictors of interest were
            nlang_z and entropy_z; covariates were age_z, edu_z, sex_binary, and mean_FD_z (mean
            framewise displacement, Power et al., z-scored). Simple Bonferroni adjustment across
            the three outcomes was applied for the two predictors of interest."""
        )
    )
    add_figure_caption(
        doc,
        "Group-mean Fisher-z resting-state connectivity (six seeds, MNI-normalised BOLD, motion-regressed; generated for this report).",
    )
    add_picture_safe(doc, GEN / "fig_group_mean_fc_matrix.png", 5.5)
    add_figure_caption(
        doc,
        "Distribution of mean framewise displacement across participants (Power et al. metric).",
    )
    add_picture_safe(doc, GEN / "fig_mean_fd_distribution.png", 5.5)
    add_figure_caption(
        doc,
        "Descriptive relationship between mean network Fisher-z and nlang (covariates are not partialled out in the plot).",
    )
    add_picture_safe(doc, GEN / "fig_mean_fc_vs_nlang.png", 5.5)

    doc.add_heading("3.5 Diffusion MRI and TBSS", level=2)
    doc.add_paragraph(
        p(
            """Per subject, eddy-corrected data and dtifit produced FA maps in
            derivatives/dwi_processed/. TBSS registered all FA images, built a mean FA skeleton, and
            projected each subject onto the skeleton (all_FA_skeletonised.nii.gz). Group inference
            used FSL randomise with 5000 permutations, TFCE, and the same five demographic regressors
            as in the design matrix (without mean FD). Contrasts tested positive and negative
            relationships of FA with nlang and with entropy."""
        )
    )

    multimodal_dir = BASE / "derivatives" / "results" / "figures"
    mm_dwi = multimodal_dir / "fig_coreg_dwi_mni.png"
    mm_fmri = multimodal_dir / "fig_coreg_fmri_mni.png"
    mm_coupling = multimodal_dir / "fig8_roi_coupling.png"
    doc.add_heading(
        "3.6 Multimodal alignment and exploratory structure–function coupling", level=2
    )
    if mm_dwi.is_file() or mm_fmri.is_file() or mm_coupling.is_file():
        doc.add_paragraph(
            p(
                """For methods documentation, both modalities were brought into standard space suitable for
                group analysis: fMRI after T1-based coregistration and spatial normalization (above);
                DWI-derived FA after FSL preprocessing and TBSS registration to the study template.
                Figures below visualize group TBSS mean FA (and skeleton, when available) on an MNI T1
                underlay, and one cohort resting-state volume after SPM normalization on the same
                template class (rendering produced by coregistration_qc_and_coupling.py; volume selected
                by a fixed, sorted directory listing of pipeline outputs—no cherry-picking). As an
                exploratory integration step not entering the primary inferential hypotheses, Pearson
                correlations were computed at each language-network ROI between mean skeletonised FA
                (per subject from all_FA_skeletonised.nii.gz, volume order from tbss_design/subject_order.txt)
                and mean language-network Fisher-z (fc_outcomes_motion.csv). This tests whether local
                white-matter anisotropy co-varies with FC strength at the same MNI coordinates across
                participants; coefficients are descriptive and uncorrected for multiple ROIs."""
            )
        )
        if mm_dwi.is_file():
            add_figure_caption(
                doc,
                "Diffusion: group mean FA (TBSS) and mean FA skeleton on MNI underlay (fig_coreg_dwi_mni.png).",
            )
            add_picture_safe(doc, mm_dwi, 6.2)
        if mm_fmri.is_file():
            add_figure_caption(
                doc,
                "fMRI: resting-state volume in MNI152 after SPM25 coregistration, normalization, and smoothing (fig_coreg_fmri_mni.png).",
            )
            add_picture_safe(doc, mm_fmri, 6.2)
        if mm_coupling.is_file():
            add_figure_caption(
                doc,
                "Exploratory ROI-wise correlation of skeleton FA with mean language FC (fig8_roi_coupling.png; companion table roi_structure_function.csv).",
            )
            add_picture_safe(doc, mm_coupling, 5.8)
    else:
        doc.add_paragraph(
            p(
                """Multimodal QC figures (mean FA/skeleton on MNI; SPM-normalised resting fMRI on MNI152) and
                the exploratory structure–function coupling panel are generated by
                coregistration_qc_and_coupling.py into derivatives/results/figures/. They were not found at
                report build time; run that script, then rebuild this document to embed fig_coreg_dwi_mni.png,
                fig_coreg_fmri_mni.png, and fig8_roi_coupling.png."""
            )
        )

    doc.add_heading("4. Results", level=1)

    doc.add_heading("4.1 Functional connectivity", level=2)
    doc.add_paragraph(
        p(
            """Across the three pre-specified outcomes, neither nlang_z nor entropy_z reached
            significance after Bonferroni correction across outcomes. Standardised coefficients
            were small. Sex (male vs female) showed lower mean connectivity and lower IFG–STG Fisher-z
            on both hemispheres at uncorrected p < 0.05. Mean FD was positively associated with
            right IFG–STG coupling (p = 0.023), consistent with motion-related variance in that edge.
            Tabulated coefficients are in derivatives/nilearn_fc_motion/fc_glm_motion_results.csv."""
        )
    )

    doc.add_heading("4.2 White-matter microstructure (TBSS)", level=2)
    doc.add_paragraph(
        p(
            """Voxelwise permutation testing on the skeleton found no voxels with TFCE
            corrected p-maps exceeding the conventional FWE threshold (corrp > 0.95) for any of the
            four contrasts. Raw t-map ranges and maximum corrp values are summarised in
            derivatives/tbss_inspection_report.txt. Thus the pre-registered structural hypothesis
            was not supported at the chosen threshold in this sample and model."""
        )
    )

    doc.add_heading("4.3 Supplemental figures from the project tree", level=2)
    doc.add_paragraph(
        p(
            """Additional panels were saved earlier in the derivatives folder; they are embedded
            below when the file path still exists on disk. Missing files are skipped with a short
            inline note so you can regenerate or restore them."""
        )
    )
    for path, caption in extra_figures:
        add_figure_caption(doc, caption)
        add_picture_safe(doc, path, width_in=6.2)
        doc.add_paragraph()

    doc.add_heading("5. Discussion", level=1)
    doc.add_paragraph(
        p(
            """Null group findings for nlang and entropy under strict correction do not imply that
            multilingual experience is irrelevant to the brain; they indicate that linear effects
            of these questionnaire scores on these particular summaries were not detectable in n=51.
            Possible factors include limited power, correlated predictors, restriction of range in an
            educated multilingual sample, summary measures (FA only; seed-based FC only), and the
            absence of task-based localisers for subject-specific ROIs. Sex and motion effects on FC
            highlight the importance of nuisance modelling in resting-state work."""
        )
    )

    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        p(
            """This multimodal workflow—from SPM normalisation through Nilearn FC and TBSS with
            randomise—yields an auditable analysis path with open data. The main inferential results
            for the primary hypotheses were null under the stated thresholds; secondary associations
            involving sex and motion warrant cautious interpretation. Future work could add NODDI or
            fixel models, task-defined ROIs (e.g. Alice localiser), or multiverse specifications."""
        )
    )

    doc.add_heading("References (selected)", level=1)
    refs = [
        "Ashburner, J., & Friston, K. J. (2005). Unified segmentation. NeuroImage, 26, 839–851.",
        "Marian, V., Blumenfeld, H. K., & Kaushanskaya, M. (2007). LEAP-Q. JSLHR, 50(4), 940–967.",
        "Power, J. D., et al. (2012). Motion and FC. NeuroImage, 59, 2142–2154.",
        "Smith, S. M., et al. (2006). TBSS. NeuroImage, 31, 1487–1505.",
        "Smith, S. M., & Nichols, T. E. (2009). TFCE. NeuroImage, 44, 83–98.",
        "NEBULA101: OpenNeuro ds005613 dataset documentation.",
    ]
    for r in refs:
        doc.add_paragraph(r)

    doc.add_heading("Appendix: Key paths", level=1)
    doc.add_paragraph("PNG files generated for this Word document live under:")
    for g in generated:
        doc.add_paragraph(str(g.relative_to(BASE)))
    doc.add_paragraph(
        "Multimodal QC and coupling figures (section 3.6) are under derivatives/results/figures/: "
        "fig_coreg_dwi_mni.png, fig_coreg_fmri_mni.png, fig8_roi_coupling.png; tab-separated "
        "roi_structure_function.csv is alongside fig8."
    )
    doc.add_paragraph(
        "[PLACEHOLDER: Institutional approval or data-use statement if required for your course or journal.]"
    )

    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


def main() -> None:
    ensure_assets()
    generated = [
        fig_mean_fc_matrix(),
        fig_cohort_predictors(),
        fig_motion_fd(),
        fig_predictor_correlation(),
        fig_fc_vs_nlang(),
    ]

    candidates = [
        (BASE / "derivatives/results/figures/fig1_predictors.png", "Pre-made: predictors / design summary."),
        (BASE / "derivatives/results/figures/fig2_pipeline.png", "Pre-made: methods pipeline schematic."),
        (BASE / "derivatives/results/figures/fig3_tbss.png", "Pre-made: TBSS / skeleton overview."),
        (BASE / "derivatives/results/figures/fig4_fc_results.png", "Pre-made: FC results summary."),
        (BASE / "derivatives/results/figures/fig5_results_table.png", "Pre-made: results table figure."),
        (BASE / "derivatives/figures/group_fc_matrix.png", "Pre-made: group FC matrix (make_fc_figures.py, nilearn_fc_motion)."),
        (BASE / "derivatives/figures/connectome_glass.png", "Pre-made: glass-brain connectome (same)."),
        (BASE / "derivatives/dwi_processed/tbss/FA/slicesdir/grota.png", "TBSS slicesdir QC montage (all subjects FA)."),
        (BASE / "derivatives/dwi_processed/tbss/FA/slicesdir/sub-pp128_FA_FA.png", "Example single-subject FA panel."),
    ]
    extra = [(p, cap) for p, cap in candidates if p.is_file()]

    build_document(generated, extra)


if __name__ == "__main__":
    main()
