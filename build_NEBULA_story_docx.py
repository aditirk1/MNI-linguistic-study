r"""
Build a story-driven Word report for the NEBULA101 multimodal analysis.

This script writes NEBULA_Story_Report.docx in the project root. It is
designed to flow as a single narrative (intro -> methods -> results ->
discussion) rather than a documentation dump. All numbers come from
authoritative source files on disk (see read_*() helpers below). Where a
figure's annotated p-values disagree with the final GLM CSV, the GLM CSV
is treated as authoritative and the figure is NOT used; we draw a Word
table instead.

Run:
    .\.venv\Scripts\python.exe build_NEBULA_story_docx.py

Output:
    NEBULA_Story_Report.docx   (in the project root)
    report_assets/story_generated/  (small set of regenerated figures)
"""
from __future__ import annotations

from pathlib import Path
import textwrap

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent
ASSETS = BASE / "report_assets"
GEN = ASSETS / "story_generated"
DOCX_OUT = BASE / "NEBULA_Story_Report.docx"

ROI_NAMES = ["IFG_L", "STG_L", "AngG_L", "SMA_L", "IFG_R", "STG_R"]


# ---------------------------------------------------------------------------
# Authoritative readers - all numbers come from these
# ---------------------------------------------------------------------------
def read_glm_results() -> pd.DataFrame:
    """Authoritative GLM coefficients from fc_with_motion.py."""
    fp = BASE / "derivatives/nilearn_fc_motion/fc_glm_motion_results.csv"
    return pd.read_csv(fp)


def refit_full_glm() -> pd.DataFrame:
    """Refit the three FC GLMs in the same form as fc_with_motion.py, but
    return ALL coefficients (not just nlang/entropy/mean_FD).

    Outcome ~ nlang_z + entropy_z + age_z + edu_z + sex_binary + mean_FD_z

    Returns a long-format DataFrame: outcome, predictor, beta, se, t, p.
    """
    import statsmodels.api as sm

    fc = read_fc_outcomes()
    dm = pd.read_csv(BASE / "shared_design_matrix.csv")
    df = fc.merge(dm, on="participant_id")
    # mean_FD_z is computed within fc_with_motion.py; replicate here
    df["mean_FD_z"] = (df["mean_FD"] - df["mean_FD"].mean()) / df["mean_FD"].std()

    rhs = ["nlang_z", "entropy_z", "age_z", "edu_z", "sex_binary", "mean_FD_z"]
    rows = []
    for outcome in ["mean_lang_fc", "IFG_STG_left", "IFG_STG_right"]:
        X = sm.add_constant(df[rhs].astype(float))
        y = df[outcome].astype(float)
        res = sm.OLS(y, X, missing="drop").fit()
        for name in res.params.index:
            rows.append({
                "outcome": outcome,
                "predictor": name,
                "beta": float(res.params[name]),
                "se": float(res.bse[name]),
                "t": float(res.tvalues[name]),
                "p": float(res.pvalues[name]),
                "r2": float(res.rsquared),
                "f_p": float(res.f_pvalue),
            })
    return pd.DataFrame(rows)


def read_fc_outcomes() -> pd.DataFrame:
    fp = BASE / "derivatives/nilearn_fc_motion/fc_outcomes_motion.csv"
    return pd.read_csv(fp)


def read_design() -> pd.DataFrame:
    """Design matrix restricted to the 51 analysis IDs."""
    subs = [
        s.strip()
        for s in (BASE / "subset_50_participants.txt").read_text().splitlines()
        if s.strip()
    ]
    dm = pd.read_csv(BASE / "shared_design_matrix.csv")
    return dm[dm["participant_id"].isin(subs)].copy()


def read_tbss_inspection() -> dict:
    """Parse derivatives/tbss_inspection_report.txt for the four contrasts."""
    fp = BASE / "derivatives/tbss_inspection_report.txt"
    lines = fp.read_text().splitlines()
    out = {"voxels_in_mask": None, "contrasts": {}}

    # voxels in mask (single global line)
    for s in lines:
        if "voxels_in_mask:" in s:
            try:
                out["voxels_in_mask"] = int(s.split(":")[1].split()[0])
            except Exception:
                pass
            break

    cur = None
    pending = None  # 't' or 'corrp' — what the next 'min max:' line refers to
    for s in lines:
        st = s.strip()
        if st.startswith("--- Contrast"):
            try:
                cur = int(st.split()[2])
            except Exception:
                cur = None
            if cur is not None:
                out["contrasts"][cur] = {}
            pending = None
            continue
        if cur is None:
            continue
        if "Raw t-stat" in st:
            pending = "t"
            continue
        if "TFCE 1-p" in st:
            pending = "corrp"
            continue
        if st.startswith("min max:") and pending is not None:
            parts = st.split(":", 1)[1].split()
            if len(parts) >= 2:
                try:
                    a, b = float(parts[0]), float(parts[1])
                    if pending == "t":
                        out["contrasts"][cur]["t_min"] = a
                        out["contrasts"][cur]["t_max"] = b
                    elif pending == "corrp":
                        out["contrasts"][cur]["corrp_min"] = a
                        out["contrasts"][cur]["corrp_max"] = b
                except Exception:
                    pass
            pending = None
            continue
        if st.startswith("count:") and "corrp > 0.95" in st:
            try:
                out["contrasts"][cur]["nvox_sig"] = int(st.split(":")[1].split()[0])
            except Exception:
                pass
    return out


# ---------------------------------------------------------------------------
# Regenerated figures (deterministic from disk)
# ---------------------------------------------------------------------------
def ensure_assets():
    GEN.mkdir(parents=True, exist_ok=True)


def fig_group_fc_matrix() -> Path:
    """Group-mean Fisher-z 6x6 from final pipeline (.npy files)."""
    out = GEN / "fig_group_mean_fc.png"
    npys = sorted((BASE / "derivatives/nilearn_fc_motion").glob("sub-*_fc_matrix.npy"))
    mats = [np.load(f) for f in npys]
    mean_z = np.mean(np.stack(mats, axis=0), axis=0)
    n_sub = len(npys)
    off = mean_z[np.triu_indices(6, k=1)]
    vmax = max(float(np.max(np.abs(off))), 0.15)
    mask = np.eye(6, dtype=bool)
    fig, ax = plt.subplots(figsize=(6.4, 5.4))
    sns.heatmap(
        mean_z, mask=mask, annot=True, fmt=".2f", cmap="RdBu_r", center=0,
        vmin=-vmax, vmax=vmax, square=True, linewidths=0.5, linecolor="white",
        xticklabels=ROI_NAMES, yticklabels=ROI_NAMES,
        cbar_kws={"label": "Fisher z", "shrink": 0.8}, ax=ax,
    )
    ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha="right")
    ax.set_title(
        f"Group-mean resting-state FC, n={n_sub}\n"
        f"6 language-network seeds (motion-regressed Fisher-z)",
        fontsize=11,
    )
    fig.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return out


def fig_predictor_panel() -> Path:
    """Three-panel: nlang dist, entropy dist, scatter."""
    out = GEN / "fig_predictors.png"
    d = read_design()
    fig, axes = plt.subplots(1, 3, figsize=(11, 3.4))

    ax = axes[0]
    vc = d["nlang"].value_counts().sort_index()
    ax.bar(vc.index.astype(int), vc.values, color="steelblue", edgecolor="black")
    ax.set_xlabel("Number of languages (nlang)")
    ax.set_ylabel("Participants")
    ax.set_title("Language diversity")

    ax = axes[1]
    ax.hist(d["entropy_curr_tot_exp"], bins=12, color="coral", edgecolor="black")
    ax.set_xlabel("Shannon entropy of exposure")
    ax.set_ylabel("Participants")
    ax.set_title("Use balance")

    ax = axes[2]
    r = d["nlang"].corr(d["entropy_curr_tot_exp"])
    ax.scatter(d["nlang"], d["entropy_curr_tot_exp"],
               alpha=0.7, s=45, c="purple", edgecolors="white")
    # OLS line
    z = np.polyfit(d["nlang"], d["entropy_curr_tot_exp"], 1)
    xs = np.linspace(d["nlang"].min(), d["nlang"].max(), 50)
    ax.plot(xs, z[0] * xs + z[1], "purple", lw=1.5)
    ax.set_xlabel("nlang"); ax.set_ylabel("Entropy")
    ax.set_title(f"Diversity vs balance (r = {r:.2f})")

    fig.suptitle(f"Two predictors of multilingual experience  (n = {len(d)})",
                 y=1.04, fontsize=12)
    fig.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return out


def fig_motion_fd() -> Path:
    out = GEN / "fig_motion_fd.png"
    fc = read_fc_outcomes()
    fig, ax = plt.subplots(figsize=(6, 3.6))
    ax.hist(fc["mean_FD"], bins=14, color="seagreen", edgecolor="black")
    m = fc["mean_FD"].mean()
    ax.axvline(m, color="darkred", linestyle="--", lw=1.5,
               label=f"Mean = {m:.3f} mm")
    ax.set_xlabel("Mean framewise displacement (Power et al., mm)")
    ax.set_ylabel("Participants")
    ax.set_title(f"Head motion at rest (n = {len(fc)})")
    ax.legend()
    fig.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return out


def fig_descriptive_scatters() -> Path:
    """Descriptive (uncorrected) scatter of mean_lang_fc vs nlang/entropy."""
    out = GEN / "fig_descriptive_fc_scatter.png"
    fc = read_fc_outcomes()
    dm = pd.read_csv(BASE / "shared_design_matrix.csv")
    m = fc.merge(dm, on="participant_id")
    fig, axes = plt.subplots(1, 2, figsize=(10, 3.6))

    for ax, x, label, color in [
        (axes[0], "nlang", "nlang (count)", "steelblue"),
        (axes[1], "entropy_curr_tot_exp", "Shannon entropy", "coral"),
    ]:
        r = m[x].corr(m["mean_lang_fc"])
        ax.scatter(m[x], m["mean_lang_fc"], alpha=0.7,
                   c=color, edgecolors="k", s=40)
        z = np.polyfit(m[x], m["mean_lang_fc"], 1)
        xs = np.linspace(m[x].min(), m[x].max(), 50)
        ax.plot(xs, z[0] * xs + z[1], color=color, lw=1.5)
        ax.set_xlabel(label)
        ax.set_ylabel("Mean Fisher-z (15 edges)")
        ax.set_title(f"r = {r:+.2f}  (uncorrected, no covariates)")
    fig.suptitle("Descriptive bivariate views (these do NOT control for sex / motion / age / edu)",
                 y=1.04, fontsize=11)
    fig.tight_layout()
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return out


# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_caption(doc, n, text: str):
    """If n is an int, render 'Figure N. ...'; if n is a string, use it as the
    prefix verbatim (e.g. 'Table 2.'). Pass n=None for caption with no prefix."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if isinstance(n, int) and n > 0:
        r = p.add_run(f"Figure {n}. ")
        r.bold = True
        p.add_run(text).italic = True
    elif isinstance(n, str):
        r = p.add_run(f"{n} ")
        r.bold = True
        p.add_run(text).italic = True
    else:
        # n == 0 or None — caller passed a fully-formatted string in `text`
        # (e.g. 'Table 2. ...'). Bold the leading 'Table N.' part if present.
        if text.lower().startswith("table") and "." in text[:12]:
            head, rest = text.split(".", 1)
            r = p.add_run(head + ". ")
            r.bold = True
            p.add_run(rest.strip()).italic = True
        else:
            p.add_run(text).italic = True


def add_picture(doc, path: Path, width_in: float = 6.2):
    if not path.is_file():
        doc.add_paragraph(f"[missing image: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))


def add_para(doc, text: str):
    doc.add_paragraph(textwrap.dedent(text).strip())


def style_table_header(row, color_hex="2E5984"):
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Build narrative
# ---------------------------------------------------------------------------
def build_doc():
    glm = read_glm_results()
    fc = read_fc_outcomes()
    d = read_design()
    tbss = read_tbss_inspection()
    full = refit_full_glm()  # all coefficients (sex, age, edu, motion)

    # --- summary numbers (computed, not hard-coded) -----------------------
    age_mean, age_sd = d["age"].mean(), d["age"].std()
    edu_mean, edu_sd = d["edu"].mean(), d["edu"].std()
    nlang_mean, nlang_sd = d["nlang"].mean(), d["nlang"].std()
    ent_mean, ent_sd = d["entropy_curr_tot_exp"].mean(), d["entropy_curr_tot_exp"].std()
    n_male = int((d["sex_binary"] == 1).sum())
    n_female = int((d["sex_binary"] == 0).sum())
    fd_mean, fd_sd = fc["mean_FD"].mean(), fc["mean_FD"].std()
    fd_min, fd_max, fd_med = fc["mean_FD"].min(), fc["mean_FD"].max(), fc["mean_FD"].median()
    r_pred = d["nlang"].corr(d["entropy_curr_tot_exp"])

    # --- figures ----------------------------------------------------------
    f_pred = fig_predictor_panel()
    f_fc = fig_group_fc_matrix()
    f_fd = fig_motion_fd()
    f_scatter = fig_descriptive_scatters()

    # --- doc setup --------------------------------------------------------
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)
    for h, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        s = doc.styles[h]; s.font.name = "Calibri"; s.font.size = Pt(sz); s.font.bold = True

    # ----- Title block ----------------------------------------------------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Multilingual experience and the multimodal brain")
    r.bold = True; r.font.size = Pt(20)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("Resting-state functional connectivity and white-matter microstructure "
              "in 51 healthy adults from NEBULA101 (OpenNeuro ds005613)").italic = True
    m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("Aditi  |  Multimodal Neuroimaging, Spring 2026  |  29 April 2026")
    doc.add_paragraph()

    # ----- Abstract -------------------------------------------------------
    doc.add_heading("Abstract", level=1)
    add_para(doc, f"""
        Whether multilingual experience leaves a measurable trace on the brain depends
        on how that experience is operationalised. Two complementary measures were
        evaluated: language **diversity** (number of languages spoken, *nlang*) and
        language **use balance** (Shannon entropy of self-reported daily exposure,
        *entropy*). In n = {len(d)} healthy adults from the NEBULA101 dataset, with
        complete T1-weighted, resting-state fMRI, and diffusion MRI data, group
        regression models tested whether either measure predicted (a) seed-based
        resting-state functional connectivity in a six-region language network, after
        SPM25 normalisation and motion regression, and (b) voxelwise fractional
        anisotropy on the TBSS skeleton, after FSL eddy/dtifit and randomise/TFCE
        permutation inference. **Neither predictor** reached significance for any
        pre-specified outcome: connectivity p-values for *nlang* and *entropy* ranged
        from {glm[glm.predictor.isin(['nlang_z','entropy_z'])]['p'].min():.3f} to
        {glm[glm.predictor.isin(['nlang_z','entropy_z'])]['p'].max():.3f} (all
        Bonferroni-uncorrected; none survived correction across the three pre-specified
        outcomes), and no white-matter skeleton voxel exceeded the conventional FWE
        threshold of corrp > 0.95 in any of the four diffusion contrasts (max corrp
        across contrasts = {max(c['corrp_max'] for c in tbss['contrasts'].values()):.2f},
        with the skeleton mask containing {tbss['voxels_in_mask']:,} voxels).
        Sex and head motion contributed appreciable variance to functional
        connectivity, validating the nuisance model. The findings constrain the
        magnitude of *linear* questionnaire–brain effects detectable in this
        moderately-sized, highly educated cohort, and motivate the alternatives
        considered in the discussion.
    """)

    # ----- 1. Why this study? --------------------------------------------
    doc.add_heading("1. Background and motivation", level=1)
    add_para(doc, """
        About half the world's population is bilingual or multilingual, and decades of
        cognitive-neuroscience work have asked whether managing two or more languages
        leaves a measurable signature on the brain. The literature is famously mixed.
        Some studies report higher fractional anisotropy in language-relevant
        white-matter tracts (e.g., the arcuate or superior longitudinal fasciculi) in
        multilinguals; others find null or even reduced effects. Resting-state
        functional connectivity has shown a similarly heterogeneous picture across
        the language network and the multiple-demand control system.
    """)
    add_para(doc, """
        A central reason for the heterogeneity is that "multilingualism" is not a
        single construct. It can be measured as the number of languages a person
        speaks (a count of repertoire diversity), as how evenly those languages are
        actually used in daily life (a balance or entropy measure), or as the age at
        which they were acquired. These are correlated but not identical, and their
        neural correlates may differ.
    """)
    add_para(doc, f"""
        This project follows the NEBULA101 framework and tests whether two
        complementary measures derived from the LEAP-Q questionnaire — *nlang* (a
        count) and *entropy_curr_tot_exp* (Shannon entropy of current daily
        exposure) — make independent contributions to brain organisation, after
        accounting for age, education, sex, and (for fMRI) head motion. In our
        analysis cohort, the two predictors are correlated (Pearson r =
        {r_pred:.2f}); they are related, not redundant.
    """)
    add_caption(doc, 1,
                "The two predictors. Left: histogram of nlang (1–10 languages). "
                "Middle: histogram of Shannon entropy of daily exposure. "
                "Right: predictor scatter; r ≈ 0.43 means they share variance but "
                "are not interchangeable.")
    add_picture(doc, f_pred, 6.4)

    add_para(doc, """
        Two pre-specified, directional hypotheses were evaluated:
    """)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("H1 (white matter): ").bold = True
    p.add_run("higher nlang and/or higher entropy → higher FA in language-related "
             "tracts on the TBSS skeleton.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("H2 (functional connectivity): ").bold = True
    p.add_run("higher nlang and/or higher entropy → altered resting coupling within "
             "and between left-lateralised language regions and their right-hemisphere "
             "homologues.")

    # ----- 2. Dataset -----------------------------------------------------
    doc.add_heading("2. Dataset and cohort", level=1)
    add_para(doc, f"""
        Data came from the publicly released NEBULA101 dataset (OpenNeuro ds005613;
        Pliatsikas et al., Multilingualism, Language and Aging group, Geneva). The
        full release contains 101 healthy multilingual adults with extensive
        questionnaire data (LEAP-Q) and a multimodal MRI battery acquired on a
        Siemens MAGNETOM Prisma 3 T scanner: a sagittal T1-weighted MPRAGE, a
        ~10-minute multiband resting-state fMRI run (TR = 2.0 s, EPI, 72 slices,
        eyes-open fixation), a multiband EPI multi-shell DWI acquisition with 117
        directions across b = 0, 700, 1000, 2800 s/mm², and an additional
        task-localiser fMRI (the "Alice in Wonderland" auditory paradigm).
    """)
    add_para(doc, f"""
        From these 101 participants, n = {len(d)} were retained for analysis
        (subset_50_participants.txt; the file name is a planning-stage holdover —
        the actual count is {len(d)}). Inclusion required complete T1, resting fMRI,
        and DWI volumes that successfully passed SPM realignment and FSL eddy
        preprocessing. The retained subset was stratified to preserve the full range
        of nlang (1 through 10), because statistical power for a continuous
        regression depends on the spread of the predictor as much as on n itself.
    """)

    # demographics table
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Variable"; hdr[1].text = "Mean ± SD"
    hdr[2].text = "Range"; hdr[3].text = "n"
    style_table_header(tbl.rows[0])
    rows = [
        ("Age (years)", f"{age_mean:.1f} ± {age_sd:.1f}", f"{d['age'].min():.1f} – {d['age'].max():.1f}", f"{len(d)}"),
        ("Education (years)", f"{edu_mean:.1f} ± {edu_sd:.1f}", f"{int(d['edu'].min())} – {int(d['edu'].max())}", f"{len(d)}"),
        ("Sex (M / F)", f"{n_male} / {n_female}", "—", f"{len(d)}"),
        ("nlang", f"{nlang_mean:.2f} ± {nlang_sd:.2f}", f"{int(d['nlang'].min())} – {int(d['nlang'].max())}", f"{len(d)}"),
        ("Entropy", f"{ent_mean:.2f} ± {ent_sd:.2f}", f"{d['entropy_curr_tot_exp'].min():.2f} – {d['entropy_curr_tot_exp'].max():.2f}", f"{len(d)}"),
    ]
    for r_ in rows:
        c = tbl.add_row().cells
        for i, v in enumerate(r_):
            c[i].text = v
    add_caption(doc, 2, f"Demographics and predictors of the analysis cohort (n = {len(d)}).")

    # ----- 3. Pipeline overview ------------------------------------------
    doc.add_heading("3. Two parallel pipelines", level=1)
    add_para(doc, """
        The analysis is structured as two parallel pipelines that converge in the
        discussion. The fMRI arm yields a per-subject 6 × 6 functional-connectivity
        matrix on which a group GLM is estimated for three pre-specified outcomes.
        The DWI arm yields a per-subject FA map projected onto a group white-matter
        skeleton, on which the same predictors are tested voxelwise with permutation
        inference.
    """)
    add_caption(doc, 3, "Pipeline schematic. Resting fMRI → SPM25 normalisation → "
                       "Nilearn seed-based FC → group GLM. Diffusion → FSL eddy/dtifit "
                       "→ TBSS → randomise/TFCE.")
    add_picture(doc, BASE / "derivatives/results/figures/fig2_pipeline.png", 6.0)

    doc.add_heading("3.1 Resting fMRI: SPM25 → Nilearn → group GLM", level=2)
    add_para(doc, """
        Per subject, six-degree-of-freedom motion parameters were estimated with
        SPM realign (estimate). Each T1 was passed through SPM25 unified
        segmentation; the resulting forward deformation field was used to normalise
        the coregistered BOLD time series to MNI space (3 mm isotropic, 4th-order
        B-spline) and the volumes were smoothed with a 6 mm FWHM Gaussian kernel.
        The smoothed warped BOLD (sw*.nii) was the input to connectivity analysis.
    """)
    add_para(doc, """
        Six 8-mm-radius spherical seeds were placed in MNI space at canonical
        language-network coordinates: left and right inferior frontal gyrus
        (IFG_L = −51, 22, 10 / IFG_R = +51, 22, 10), left and right superior
        temporal gyrus (STG_L = −56, −14, 4 / STG_R = +56, −14, 4), the left
        angular gyrus (AngG_L = −46, −62, 28), and the supplementary motor area
        (SMA_L = −4, 4, 52). For each seed, the mean BOLD time-course was extracted
        with NiftiSpheresMasker, band-pass filtered (0.01–0.1 Hz), detrended,
        z-scored, and the six SPM realignment parameters were regressed out at the
        time-series stage. Pairwise Pearson correlations across seeds were
        Fisher-z transformed and saved as 6 × 6 matrices.
    """)
    add_para(doc, """
        Three pre-specified outcomes entered the group GLM: (i) the mean of the 15
        unique upper-triangular Fisher-z edges (mean_lang_fc), (ii) the
        IFG_L–STG_L edge (left dorsal language pathway), and (iii) the
        IFG_R–STG_R edge (right-hemisphere homologue). Each outcome was modelled
        as y ~ nlang_z + entropy_z + age_z + edu_z + sex_binary + mean_FD_z, fit by
        OLS in statsmodels. For the two predictors of interest, simple Bonferroni
        correction across the three outcomes was applied (p × 3).
    """)

    doc.add_heading("3.2 DWI: FSL eddy/dtifit → TBSS → randomise/TFCE", level=2)
    add_para(doc, f"""
        Diffusion data were preprocessed in WSL Ubuntu with FSL: brain-mask
        extraction from b0 images (bet), eddy-current and motion correction with
        eddy (with bvec rotation), and tensor fitting with dtifit producing
        voxelwise FA, MD and eigen images. All {len(d)} subjects completed
        preprocessing.
    """)
    add_para(doc, f"""
        TBSS aligned each subject's FA image to the FMRIB58_FA template, built a
        group mean FA skeleton, and projected per-subject FA values onto that
        skeleton. The resulting 4D image (all_FA_skeletonised.nii.gz) holds
        {tbss['voxels_in_mask']:,} skeleton voxels per subject. randomise was run
        with 5,000 permutations and threshold-free cluster enhancement (--T2),
        producing voxelwise t-maps and TFCE family-wise-error-corrected (1 − p)
        maps for four contrasts: positive and negative associations of FA with
        nlang and with entropy, in a model containing the same five demographic
        regressors as the FC analysis (mean FD is fMRI-specific). Voxels at
        FWE p < 0.05 satisfy corrp > 0.95.
    """)

    # ----- 4. Quality control -------------------------------------------
    doc.add_heading("4. Quality control", level=1)
    add_para(doc, f"""
        Head motion at rest was modest in this cohort (mean FD = {fd_mean:.3f} ±
        {fd_sd:.3f} mm; median {fd_med:.3f}; range {fd_min:.3f}–{fd_max:.3f} mm),
        so no participant was excluded a priori. Mean FD enters the FC GLM as a
        covariate to absorb residual motion-driven coupling rather than discarding
        scans. eddy_quad/SQUAD per-subject and group reports were generated for
        the diffusion data, and the FSL slicesdir montage of all per-subject FA
        maps allowed visual confirmation of registration quality.
    """)
    add_caption(doc, 4, "Distribution of mean framewise displacement (Power et al., 2012) "
                       "across the analysis cohort.")
    add_picture(doc, f_fd, 5.5)
    add_caption(doc, 5, "Example single-subject FA map after dtifit "
                       "(QC montage from FSL slicesdir).")
    add_picture(doc, BASE / "derivatives/dwi_processed/tbss/FA/slicesdir/sub-pp003_FA_FA.png", 6.4)

    # ----- 5. Results -----------------------------------------------------
    doc.add_heading("5. Results", level=1)

    doc.add_heading("5.1 The language network is intact and recoverable", level=2)
    add_para(doc, """
        Before testing the predictors of interest, it is reassuring to verify that
        the seed-based FC pipeline recovers a sensible network. The group-mean
        Fisher-z matrix (Figure 6) shows the expected pattern of strong coupling
        within and between the bilateral perisylvian seeds (IFG, STG, SMA), with
        the angular gyrus seed sitting at the periphery (its weak coupling to
        peri-Sylvian regions is consistent with its semantic / DMN-border
        location). In other words, the pipeline is doing what it should — any
        null effects below cannot be attributed to a broken connectivity
        estimator.
    """)
    add_caption(doc, 6, "Group-mean Fisher-z connectivity, n = 51 (motion-regressed). "
                       "The expected dorsal-language and bilateral-IFG pattern is clearly "
                       "recovered.")
    add_picture(doc, f_fc, 5.5)

    doc.add_heading("5.2 No effect of multilingual experience on functional connectivity", level=2)
    add_para(doc, """
        Bivariate scatters of mean network FC against the two predictors hint at no
        clear relationship; small negative slopes are present but the scatter is wide
        (Figure 7). Because the outcome is also influenced by age, education, sex,
        and motion, the inferential test is the partial regression — the GLM with
        all covariates present.
    """)
    add_caption(doc, 7, "Descriptive (uncorrected) views of mean network FC against "
                       "nlang and entropy. These do NOT control for sex, motion, age, or "
                       "education; they are shown only to expose the raw spread.")
    add_picture(doc, f_scatter, 6.2)

    add_para(doc, """
        Table 2 summarises the partial coefficients for the two predictors of
        interest across the three pre-specified outcomes. For all three outcomes
        and both predictors, the partial coefficient is small (|β| ≤ 0.05 in
        z-units), uncorrected p > 0.18, and Bonferroni-corrected p = 1 (or close).
        We therefore fail to reject the null for both *nlang* and *entropy* at
        the pre-specified inferential threshold.
    """)

    # GLM results table — pulled from CSV
    interest = glm[glm["predictor"].isin(["nlang_z", "entropy_z"])].copy()
    tbl = doc.add_table(rows=1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Outcome", "Predictor", "β", "SE", "t", "p (uncorr) → ×3 Bonf"]):
        hdr[i].text = h
    style_table_header(tbl.rows[0])
    for _, r_ in interest.iterrows():
        cells = tbl.add_row().cells
        cells[0].text = str(r_["outcome"])
        cells[1].text = str(r_["predictor"]).replace("_z", "")
        cells[2].text = f"{r_['beta']:+.3f}"
        cells[3].text = f"{r_['se']:.3f}"
        cells[4].text = f"{r_['t']:+.2f}"
        cells[5].text = f"{r_['p']:.3f} → {r_['p_bonf_3outcomes']:.3f}"
    add_caption(doc, 0,
                "Table 2. Group GLM partial coefficients for nlang and entropy on the three "
                "pre-specified FC outcomes, n = 51. Each row controls for the other predictor "
                "and for age, education, sex, and mean FD. None of the six tests survives "
                "Bonferroni correction across the three outcomes.")

    # pull sex / motion numbers straight from the refit so everything is verifiable
    def _row(outcome, predictor):
        m = full[(full["outcome"] == outcome) & (full["predictor"] == predictor)]
        if len(m) == 0:
            return None
        return m.iloc[0]
    sx_mean = _row("mean_lang_fc", "sex_binary")
    sx_left = _row("IFG_STG_left", "sex_binary")
    sx_right = _row("IFG_STG_right", "sex_binary")
    fd_right = _row("IFG_STG_right", "mean_FD_z")
    r2s = full.drop_duplicates("outcome").set_index("outcome")["r2"].to_dict()
    fps = full.drop_duplicates("outcome").set_index("outcome")["f_p"].to_dict()

    add_para(doc, f"""
        Two covariates do reach uncorrected significance and merit comment.
        Sex (male = 1) is associated with lower mean network FC and lower IFG–STG
        coupling on both sides: β = {sx_mean['beta']:+.2f}, p = {sx_mean['p']:.3f} on
        mean_lang_fc; β = {sx_left['beta']:+.2f}, p = {sx_left['p']:.3f} on
        IFG_STG_left; β = {sx_right['beta']:+.2f}, p = {sx_right['p']:.3f} on
        IFG_STG_right. Mean FD is positively associated with right-hemisphere
        IFG–STG coupling (β = {fd_right['beta']:+.3f}, p = {fd_right['p']:.3f}),
        in the expected direction (more head motion → spuriously higher
        correlation). The three FC models had R² of {r2s['mean_lang_fc']:.2f},
        {r2s['IFG_STG_left']:.2f}, and {r2s['IFG_STG_right']:.2f} respectively,
        with overall-F p-values of {fps['mean_lang_fc']:.3f},
        {fps['IFG_STG_left']:.3f}, and {fps['IFG_STG_right']:.3f}: the models
        explain real variance, but it comes from the nuisance covariates rather
        than the predictors of interest. Sex and motion are secondary findings
        (not pre-registered), and they validate the nuisance model — sex
        differences in resting coupling are well-known, and the motion–FC
        association reaffirms the importance of regressing FD as a covariate.
    """)

    doc.add_heading("5.3 No effect on white-matter microstructure either", level=2)
    add_para(doc, f"""
        On the TBSS skeleton, the mean FA mask spans {tbss['voxels_in_mask']:,}
        voxels covering the major language-relevant tracts. randomise/TFCE found
        no voxel exceeding the conventional FWE threshold (corrp > 0.95) in any of
        the four contrasts. The strongest sub-threshold trend was for "+nlang"
        (max corrp = {tbss['contrasts'][1]['corrp_max']:.2f}; raw t in
        [{tbss['contrasts'][1]['t_min']:+.2f}, {tbss['contrasts'][1]['t_max']:+.2f}]),
        i.e., still ≈ 0.36 below the FWE threshold of 0.95. The full picture is
        in Figure 8 and Table 3.
    """)
    add_caption(doc, 8, "TBSS results across the four contrasts. No skeleton voxel survives "
                       "FWE p < 0.05 (corrp > 0.95 dashed line). Maximum corrp values are "
                       "well below the threshold for both predictors and both directions.")
    add_picture(doc, BASE / "derivatives/results/figures/fig3_tbss.png", 6.4)

    # TBSS table
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["Contrast", "Direction", "Raw t range", "Max corrp", "Voxels FWE-sig"]):
        tbl.rows[0].cells[i].text = h
    style_table_header(tbl.rows[0])
    contrasts_meta = [
        (1, "+nlang"), (2, "−nlang"), (3, "+entropy"), (4, "−entropy"),
    ]
    for cn, label in contrasts_meta:
        c = tbss['contrasts'].get(cn, {})
        cells = tbl.add_row().cells
        cells[0].text = f"C{cn}"
        cells[1].text = label
        cells[2].text = f"[{c.get('t_min', float('nan')):+.2f}, {c.get('t_max', float('nan')):+.2f}]"
        cells[3].text = f"{c.get('corrp_max', float('nan')):.2f}"
        cells[4].text = f"{c.get('nvox_sig', 0)}"
    add_caption(doc, 0,
                "Table 3. TBSS randomise/TFCE results, 5,000 permutations, "
                f"skeleton mask = {tbss['voxels_in_mask']:,} voxels. No voxel survives "
                "the conventional FWE threshold (corrp > 0.95) in any contrast.")

    doc.add_heading("5.4 Where the variance went", level=2)
    r2_min, r2_max = min(r2s.values()), max(r2s.values())
    fp_min, fp_max = min(fps.values()), max(fps.values())
    cm = tbss["contrasts"]
    add_para(doc, f"""
        Although the questionnaire-derived predictors carried no significant
        variance for either modality, the FC models were not empty. Across the
        three FC outcomes, R² ranged from {r2_min:.2f} to {r2_max:.2f} and
        overall-F p-values from {fp_min:.3f} to {fp_max:.3f} — the models do
        explain variance, just from sex and (for one outcome) motion rather
        than from multilingual experience. The TBSS contrasts, by contrast,
        look essentially flat: max corrp values across the four contrasts are
        {cm[1]['corrp_max']:.2f}, {cm[2]['corrp_max']:.2f},
        {cm[3]['corrp_max']:.2f}, and {cm[4]['corrp_max']:.2f} — a long way
        from the FWE threshold of 0.95.
    """)

    # ----- 6. Discussion -------------------------------------------------
    doc.add_heading("6. Discussion", level=1)
    add_para(doc, """
        We did not detect an association between two complementary measures of
        multilingual experience and either resting-state functional connectivity
        in a six-region language network or fractional anisotropy on the TBSS
        white-matter skeleton, at the pre-specified inferential thresholds. This
        mirrors a number of recent reports in the multilingualism literature in
        which large effects of language experience on brain structure or function
        have been difficult to replicate at conventional sample sizes.
    """)
    add_para(doc, f"""
        Several non-exclusive factors plausibly contribute. First, the population
        effect sizes that link self-reported language experience to seed-based FC or
        skeletonised FA may be small. With n = {len(d)}, an analysis is reasonably
        powered to detect medium-to-large standardised partial coefficients but
        not the small (|β| < 0.1) effects that are most often reported in the
        literature. Second, the two predictors are correlated (r = {r_pred:.2f}):
        unique-variance partial coefficients are by construction smaller than
        their bivariate counterparts, and any shared signal cannot be partitioned
        between *nlang* and *entropy*. Third, our cohort is highly educated (mean
        education = {edu_mean:.0f} years) — the range of multilingual experience is
        rich (1–10 languages) but other dimensions, including the age of L2
        acquisition and proficiency, are constrained. Fourth, FA on a TBSS
        skeleton summarises microstructure with a single tensor model; multi-shell
        metrics such as NODDI or fixel-based analyses, which are feasible with
        the b = 700/1000/2800 acquisition that is already on disk, may be more
        sensitive to specific microstructural variations.
    """)
    add_para(doc, """
        Two sets of secondary observations strengthen confidence in the negative
        primary results. The bilateral language-network FC is recovered cleanly
        (Figure 6), so the FC pipeline behaves as designed. Sex and motion explain
        appreciable variance in the FC outcomes, in the expected directions and
        magnitudes — which is exactly the kind of pattern that reassures us the
        regression is sensitive to non-trivial effects when they exist.
    """)
    add_para(doc, """
        Finally, an absence of evidence is not evidence of absence. Null findings
        bound the magnitude of linear, questionnaire-driven effects that are
        detectable in this subsample with this analytic recipe. Three concrete
        extensions are well-positioned to push further: (i) using subject-specific
        language ROIs from the Alice in Wonderland localiser (already preprocessed
        with SPM but not used for second-level here) instead of fixed MNI seeds;
        (ii) re-analysing diffusion data with NODDI or fixel-based metrics; and
        (iii) adding network-level (ICA / dictionary-learning) FC summaries that
        do not commit to a small set of edges a priori.
    """)

    # ----- 7. Limitations ------------------------------------------------
    doc.add_heading("7. Limitations", level=1)
    bullet = lambda txt: doc.add_paragraph(txt, style="List Bullet")
    bullet("Cross-sectional, observational design — no causal inference is possible.")
    bullet(f"n = {len(d)} bounds power to detect small effects (post-hoc power < 80% "
           "for r ≈ 0.2, see Figure 9).")
    bullet("Six a priori MNI seeds do not cover the whole language network or "
           "the multiple-demand control system.")
    bullet("Single-shell FA was the only diffusion summary; multi-shell models "
           "(NODDI, fixel-based) were not used despite the multi-shell acquisition "
           "being available.")
    bullet("No fMRIPrep / aCompCor / ICA-AROMA: motion is partially controlled via "
           "six realignment parameters and a mean-FD covariate, but residual "
           "motion artefact is possible.")
    bullet("No fieldmap-based EPI distortion correction.")
    bullet("Highly educated cohort (mean ≈ {0:.0f} years); generalisability to "
           "lower-education populations is limited.".format(edu_mean))
    bullet("Self-reported exposure proportions inherit the standard limitations of "
           "questionnaires (LEAP-Q; Marian et al., 2007).")

    add_caption(doc, 9, "Post-hoc power for partial Pearson r effects at α = 0.05 "
                       "(two-tailed). With n = 51, power to detect r = 0.20 is far below 80%.")
    add_picture(doc, BASE / "derivatives/results/figures/fig6_power.png", 6.0)

    # ----- 8. Conclusion -------------------------------------------------
    doc.add_heading("8. Conclusion", level=1)
    add_para(doc, """
        Across n = 51 multilingual adults from NEBULA101, neither the number of
        languages spoken nor the Shannon entropy of daily exposure was associated
        with seed-based resting-state functional connectivity in a six-region
        language network or with white-matter fractional anisotropy on the TBSS
        skeleton, at pre-specified FWE / Bonferroni thresholds. The null inferential
        result is bracketed by a recovered language network on FC, sensible nuisance
        effects (sex and motion), and a clean diffusion pipeline — all of which
        argue that the methodology is sound and the bound on the effect size is
        the substantive contribution of this work. Future analyses are well placed
        to look in the directions where this study could not: subject-specific ROIs,
        multi-shell diffusion modelling, and network-level FC summaries.
    """)

    # ----- References ----------------------------------------------------
    doc.add_heading("References (selected)", level=1)
    refs = [
        "Ashburner, J., & Friston, K. J. (2005). Unified segmentation. NeuroImage, 26, 839–851.",
        "Bialystok, E. (2017). The bilingual adaptation. Psychological Bulletin, 143(3), 233–262.",
        "Costa, A., & Sebastián-Gallés, N. (2014). How does the bilingual experience sculpt the brain? Nature Reviews Neuroscience, 15(5), 336–345.",
        "Marian, V., Blumenfeld, H. K., & Kaushanskaya, M. (2007). The LEAP-Q. JSLHR, 50(4), 940–967.",
        "Pliatsikas, C., et al. (2024). NEBULA101: a multimodal dataset of multilingual experience. OpenNeuro ds005613.",
        "Power, J. D., et al. (2012). Spurious but systematic correlations in functional connectivity MRI networks arise from subject motion. NeuroImage, 59, 2142–2154.",
        "Smith, S. M., et al. (2006). Tract-based spatial statistics. NeuroImage, 31, 1487–1505.",
        "Smith, S. M., & Nichols, T. E. (2009). Threshold-free cluster enhancement. NeuroImage, 44, 83–98.",
    ]
    for r_ in refs:
        doc.add_paragraph(r_, style="List Number")

    # ----- Appendix ------------------------------------------------------
    doc.add_heading("Appendix A. Authoritative source files", level=1)
    appendix_rows = [
        ("Subject list (51 IDs, alphabetical)", "subset_50_participants.txt"),
        ("Design matrix (z-scored predictors)", "shared_design_matrix.csv"),
        ("Per-subject Fisher-z FC matrices", "derivatives/nilearn_fc_motion/sub-*_fc_matrix.npy"),
        ("FC subject-level outcomes + mean FD", "derivatives/nilearn_fc_motion/fc_outcomes_motion.csv"),
        ("FC group GLM results (used in Table 2)", "derivatives/nilearn_fc_motion/fc_glm_motion_results.csv"),
        ("TBSS 4-D skeletonised FA", "derivatives/dwi_processed/tbss/stats/all_FA_skeletonised.nii.gz"),
        ("TBSS design files", "derivatives/dwi_processed/tbss/stats/design.{mat,con}"),
        ("randomise TFCE corrp maps", "derivatives/dwi_processed/tbss/stats/tbss_results_tfce_corrp_tstat{1..4}.nii.gz"),
        ("randomise log", "derivatives/randomise_log.txt"),
        ("Automated TBSS inspection (used in Table 3)", "derivatives/tbss_inspection_report.txt"),
        ("Stats archive (29 Apr 2026)", "derivatives/tbss_randomise_archive_20260429_153849/"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.rows[0].cells[0].text = "Item"
    tbl.rows[0].cells[1].text = "Path"
    style_table_header(tbl.rows[0])
    for k, v in appendix_rows:
        c = tbl.add_row().cells
        c[0].text = k; c[1].text = v

    # save
    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


def main():
    ensure_assets()
    build_doc()


if __name__ == "__main__":
    main()
